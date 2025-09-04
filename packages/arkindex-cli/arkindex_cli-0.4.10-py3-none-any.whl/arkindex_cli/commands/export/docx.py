import logging
from pathlib import Path
from uuid import UUID

from arkindex_cli.commands.export.utils import MANUAL_SOURCE, uuid_or_manual
from arkindex_export import Element, Transcription, WorkerRun, open_database
from arkindex_export.queries import list_children
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def add_docx_parser(subcommands):
    docx_parser = subcommands.add_parser(
        "docx",
        description="Read data from an exported database and generate a DOCX file.",
        help="Generates a DOCX file from an Arkindex export.",
    )
    folders = docx_parser.add_mutually_exclusive_group()
    folders.add_argument(
        "--folder-ids",
        type=UUID,
        help="IDs of folders to export.",
        nargs="+",
    )
    folders.add_argument(
        "--folder-type",
        type=str,
        help="Type of the folder elements containing the elements to be exported. Defaults to 'folder'.",
        default="folder",
    )
    docx_parser.add_argument(
        "--element-type",
        type=str,
        help="Type (slug) of the elements to export. Defaults to 'page'.",
        default="page",
    )
    docx_parser.add_argument(
        "--line-type",
        type=str,
        help="""
            Type (slug) of the elements containing transcriptions. If unspecified, transcriptions are retrieved
            from --element-type elements.
        """,
    )
    sources = docx_parser.add_mutually_exclusive_group()
    sources.add_argument(
        "--worker-run-id",
        type=uuid_or_manual,
        help=f"""
            Preferable to worker-version-id.
            '{MANUAL_SOURCE}' or UUIDs of the worker run(s) that produced the transcriptions to be exported. The order in which
            the worker runs are given as argument acts as a preference order: if there are transcriptions from multiple
            worker runs on an element, the ones from the worker run at the earliest position in this list of UUIDs will
            be exported.
        """,
        nargs="+",
        default=[],
    )
    sources.add_argument(
        "--worker-version-id",
        type=uuid_or_manual,
        help=f"""
            '{MANUAL_SOURCE}' or UUIDs of the worker version(s) that produced the transcriptions to be exported. The order in which
            the worker versions are given as argument acts as a preference order: if there are transcriptions from multiple
            worker versions on an element, the ones from the worker version at the earliest position in this list of UUIDs
            will be exported.
        """,
        nargs="+",
        default=[],
    )
    docx_parser.add_argument(
        "--merge",
        action="store_true",
        help="Merge the transcriptions from all the exported elements in one folder into a single DOCX file.",
    )
    docx_parser.add_argument(
        "-o",
        "--output",
        default=Path.cwd() / "docx",
        type=Path,
        help="Path to a folder where the outputted DOCX files will be saved. Defaults to '<current_directory>/docx/'.",
        dest="output_path",
    )

    docx_parser.set_defaults(func=run)


def get_transcription_query(element, worker_version_ids, worker_run_ids):
    """
    Yields one query per worker version / worker run id, or just one query if there are none
    """
    query = Transcription.select().where(Transcription.element_id == element.id)
    for run_id in worker_run_ids:
        if run_id == MANUAL_SOURCE:
            transcription = query.where(Transcription.worker_run_id.is_null()).limit(2)
            yield transcription
        else:
            transcription = query.where(Transcription.worker_run_id == run_id).limit(2)
            yield transcription
    for version_id in worker_version_ids:
        if version_id == MANUAL_SOURCE:
            transcription = query.where(Transcription.worker_run_id.is_null()).limit(2)
            yield transcription
        else:
            transcription = (
                query.join(WorkerRun)
                .where(WorkerRun.worker_version_id == version_id)
                .limit(2)
            )
            yield transcription
    if not len(worker_run_ids) and not len(worker_version_ids):
        transcription = query
        yield transcription


def get_transcription(worker_version_ids, worker_run_ids, element):
    """
    Retrieve one transcription for the given element, according to the preferred order
    of specified worker run or worker version IDs.
    """
    queries = get_transcription_query(element, worker_version_ids, worker_run_ids)
    # Go through the queries until a transcription is found
    for query in queries:
        transcription = list(query.limit(2))
        if len(transcription) > 1:
            raise Exception(
                f"Multiple transcriptions returned for {element.type} {element.id}."
            )
        elif len(transcription) == 1:
            return transcription[0].text


def add_docx_heading(document, title):
    title = document.add_heading(title, 3)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a line break after the heading
    document.add_paragraph("")


def generate_docx(transcription, path):
    if not len(transcription):
        return 0

    document = Document()

    for index, [heading, paragraph] in enumerate(transcription):
        add_docx_heading(document, heading)
        document.add_paragraph(paragraph)
        # Add a line break between pages
        if index < len(transcription) - 1:
            document.add_paragraph("")

    # Create any missing folder from the hierarchy
    path.parent.mkdir(parents=True, exist_ok=True)

    # Save the DOCX file
    document.save(f"{path}.docx")

    return 1


def run(
    database_path: Path,
    output_path: Path,
    profile_slug: str | None = None,
    gitlab_secure_file: Path | None = None,
    folder_ids: list[UUID] | None = [],
    folder_type: str | None = "folder",
    element_type: str | None = "page",
    line_type: str | None = None,
    worker_version_id: list[str] | None = [],
    worker_run_id: list[str] | None = [],
    merge: bool | None = False,
):
    open_database(database_path)

    # If not folder ids are specified, use all folders in the DB
    if not folder_ids:
        folder_ids = [e.id for e in Element.select().where(Element.type == folder_type)]

    saved_transcriptions_count = 0

    for folder_id in folder_ids:
        # Store folder level transcriptions (only actually used if `merge` is True)
        folder_transcriptions = []

        for element in list_children(folder_id).where(Element.type == element_type):
            if line_type:
                element_lines = []
                lines = list_children(element.id).where(Element.type == line_type)
                for line in lines:
                    line_transcription = get_transcription(
                        worker_version_id, worker_run_id, line
                    )
                    if line_transcription:
                        element_lines.append(line_transcription)
                element_transcription = [element.name, "\n".join(element_lines)]
            else:
                element_transcription = [
                    element.name,
                    get_transcription(worker_version_id, worker_run_id, element),
                ]

            # Skip generation if no transcription is found
            if not element_transcription[1]:
                continue

            if merge:
                # Populate folder_transcriptions
                folder_transcriptions.append(element_transcription)

            else:
                # Generate one file per element that has a transcription
                saved_transcriptions_count += generate_docx(
                    [element_transcription],
                    output_path / str(folder_id) / element.name,
                )

        # If `merge` is True and there are folder level transcriptions, generate one file per folder
        if merge and len(folder_transcriptions):
            generate_docx(folder_transcriptions, output_path / str(folder_id))
            saved_transcriptions_count += 1

    if saved_transcriptions_count:
        logger.info(f"Files saved to {output_path}.")
    else:
        logger.warning("No transcriptions were found.")
