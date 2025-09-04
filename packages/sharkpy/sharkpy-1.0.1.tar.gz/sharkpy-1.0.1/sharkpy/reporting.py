import os
import tempfile
import atexit
from datetime import datetime
import numpy as np
import pandas as pd
import matplotlib
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import KFold, cross_validate
from sklearn.metrics import (
    make_scorer, r2_score, mean_absolute_error, mean_squared_error, 
    accuracy_score, precision_score, recall_score, f1_score,
    roc_auc_score
)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert
from typing import Optional, Tuple, List, Dict, Any

# Set style for better looking plots
sns.set_style("whitegrid")
plt.rcParams['figure.figsize'] = (10, 6)
plt.rcParams['font.size'] = 12

def _create_temp_plot(shark: Any, kind: str, width: int = 8, height: int = 6) -> Optional[str]:
    """Create a temporary plot and return its path"""
    try:
        # Create temporary file
        fd, temp_path = tempfile.mkstemp(suffix='.png')
        os.close(fd)
        
        # Store original backend
        original_backend = matplotlib.get_backend()
        
        try:
            # Set Agg backend for plot generation
            matplotlib.use('Agg')
            
            # Generate plot
            from .plotting import plot_model
            plot_model(shark.model, shark.features, shark.target, 
                      kind=kind, show=False, save_path=temp_path)
            
        finally:
            # Restore original backend
            matplotlib.use(original_backend)
        
        return temp_path
        
    except Exception as e:
        print(f"⚠️ Could not create {kind} plot: {str(e)}")
        return None

def _get_feature_importance_section(shark: Any) -> Tuple[List[str], Optional[pd.DataFrame]]:
    """Generate feature importance section for the report and a DataFrame for DOCX table"""
    lines = ["\n📊 Feature Importance:"]
    feature_data = []
    
    if hasattr(shark.model, 'feature_importances_'):
        importances = shark.model.feature_importances_
        indices = np.argsort(importances)[::-1]
        for idx in indices:
            feature = shark.feature_names[idx]
            importance = importances[idx]
            lines.append(f" - {feature}: {importance:.3f}")
            feature_data.append({'Feature': feature, 'Importance': importance})
    elif hasattr(shark.model, 'coef_'):
        coefs = shark.model.coef_
        if len(coefs.shape) > 1:  # Multi-class logistic regression
            coefs = np.abs(coefs).mean(axis=0)
        indices = np.argsort(np.abs(coefs))[::-1]
        for idx in indices:
            feature = shark.feature_names[idx]
            coef = coefs[idx]
            lines.append(f" - {feature}: {coef:.3f} (abs. coefficient)")
            feature_data.append({'Feature': feature, 'Coefficient': coef})
    else:
        lines.append(" - Feature importance not available for this model.")
        feature_data = None
    
    feature_df = pd.DataFrame(feature_data) if feature_data else None
    return lines, feature_df

def _get_statistical_details_section(shark: Any) -> List[str]:
    """Generate statistical details section for the report"""
    lines = ["\n📈 Statistical Details:"]
    lines.append(str(shark.statistical_summary))
    return lines

def _add_table_to_doc(doc: Document, df: pd.DataFrame, title: str):
    """Helper function to add a pandas DataFrame as a table to DOCX"""
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Add header
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
    
    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = f"{value:.3f}" if isinstance(value, float) else str(value)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(12)

def _export_docx_report(path: str, shark: Any, cv_metrics_df: pd.DataFrame, train_metrics_df: pd.DataFrame, problem_type: str):
    """Export report as Word document with enhanced formatting, tables, and plots."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('🦈 SharkPy Model Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    
    # Add metadata
    doc.add_heading('Model Summary', level=1)
    metadata = [
        f"Project: {shark.project_name}",
        f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Model Type: {type(shark.model).__name__}",
        f"Problem Type: {shark.problem_type.capitalize()}",
        f"Target Variable: {shark.target.name}",
        f"Features: {', '.join(shark.feature_names)}",
        f"Training Set Shape: {shark.features.shape}"
    ]
    for line in metadata:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(12)
    
    # Add cross-validation metrics table
    _add_table_to_doc(doc, cv_metrics_df, 'Cross-Validation Metrics')
    
    # Add training metrics table
    _add_table_to_doc(doc, train_metrics_df, 'Training Set Metrics')
    
    # Add feature importance table
    _, feature_df = _get_feature_importance_section(shark)
    if feature_df is not None:
        _add_table_to_doc(doc, feature_df, 'Feature Importance')
    
    # Add statistical details (for regression)
    if hasattr(shark, 'stats_model') and shark.p_values is not None:
        doc.add_heading('Statistical Details', level=1)
        p = doc.add_paragraph(str(shark.statistical_summary))
        p.runs[0].font.size = Pt(10)
    
    # Add visualizations
    doc.add_heading('Visualizations', level=1)
    plots = [('Feature Importance', 'feature_importance')]
    if problem_type == "regression":
        plots += [
            ('Prediction Plot', 'prediction'),
            ('Residuals Plot', 'residuals')
        ]
    else:  # classification
        plots += [('Confusion Matrix', 'confusion_matrix')]
        if len(np.unique(shark.target)) == 2:  # Binary classification
            plots += [
                ('ROC Curve', 'roc'),
                ('Precision-Recall Curve', 'pr_curve'),
                ('Probability Histogram', 'proba_hist')
            ]
    
    for title, kind in plots:
        temp_path = _create_temp_plot(shark, kind)
        if temp_path and os.path.exists(temp_path):
            doc.add_heading(title, level=2)
            doc.add_picture(temp_path, width=Inches(5.5))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            os.remove(temp_path)
    
    # Save document
    doc.save(path)

def _export_txt_report(path: str, lines: List[str]):
    """Export report as text file"""
    with open(path, 'w') as f:
        for line in lines:
            f.write(line + '\n')

def _convert_docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    """Convert DOCX to PDF using available tools"""
    try:
        # Try docx2pdf first (requires MS Word)
        convert(docx_path, pdf_path)
    except Exception as e:
        try:
            # Fallback to LibreOffice if available
            import subprocess
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', 
                           '--outdir', os.path.dirname(pdf_path), docx_path], 
                           check=True, capture_output=True)
        except Exception as e2:
            raise Exception(f"Could not convert to PDF. Please install Microsoft Word or LibreOffice.\nError: {str(e2)}")

def report(self, cv_folds: int = 5, export_path: Optional[str] = None, format: str = 'txt') -> Tuple[Dict[str, np.ndarray], Dict[str, float]]:
    """
    Generate comprehensive model performance report including cross-validation.
    
    Parameters
    ----------
    self : Any
        The Shark instance
    cv_folds : int, optional
        Number of folds for K-Fold cross-validation (default: 5)
    export_path : str, optional
        Path to export the report. If None, report is only printed.
        If a directory is provided, a timestamped file will be created.
    format : str, optional
        Export format: 'txt', 'pdf', or 'docx' (default: 'txt')
    
    Returns
    -------
    tuple
        - cv_results : dict
            Dictionary containing cross-validation results
        - train_metrics : dict
            Dictionary containing training set metrics
    
    Notes
    -----
    - For PDF export, ensure Microsoft Word or LibreOffice is installed for docx2pdf conversion.
    - Visualizations include feature importance, predictions/residuals (regression), or confusion matrix/ROC/PR curves (classification).
    """
    if not hasattr(self, 'model'):
        raise ValueError("🦈 No model has been trained yet. Call learn() first.")

    # Store original backend
    original_backend = matplotlib.get_backend()
    
    try:
        # Initialize report lines
        report_lines = []
        report_lines.append("🦈 SharkPy Model Report 🦈")
        report_lines.append(f"Project: {self.project_name}")
        report_lines.append(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report_lines.append(f"Model Type: {type(self.model).__name__}")
        report_lines.append(f"Problem Type: {self.problem_type.capitalize()}")
        report_lines.append(f"Target Variable: {self.target.name}")
        report_lines.append(f"Features: {', '.join(self.feature_names)}")
        report_lines.append(f"Training Set Shape: {self.features.shape}")

        # Define scoring metrics
        if self.problem_type == "regression":
            scoring = {
                'r2': 'r2',
                'mae': make_scorer(mean_absolute_error),
                'mse': make_scorer(mean_squared_error),
                'rmse': make_scorer(lambda y, y_pred: np.sqrt(mean_squared_error(y, y_pred))),
            }
        else:
            scoring = {
                'accuracy': make_scorer(accuracy_score),
                'precision': make_scorer(precision_score, average='weighted', zero_division=0),
                'recall': make_scorer(recall_score, average='weighted', zero_division=0),
                'f1': make_scorer(f1_score, average='weighted', zero_division=0),
            }
            if hasattr(self.model, 'predict_proba'):
                scoring['roc_auc'] = make_scorer(roc_auc_score, needs_proba=True, multi_class='ovr')

        # Cross-validation
        kf = KFold(n_splits=cv_folds, shuffle=True, random_state=42)
        cv_results = cross_validate(
            self.model, 
            self.features, 
            self.target, 
            cv=kf,
            scoring=scoring,
            n_jobs=-1
        )

        # Format CV results for text and table
        report_lines.append("\n📊 Cross-Validation Metrics:")
        cv_metrics_list = []
        if self.problem_type == "regression":
            metrics = ['r2', 'mae', 'mse', 'rmse']
            for metric in metrics:
                mean = cv_results[f'test_{metric}'].mean()
                std = cv_results[f'test_{metric}'].std()
                report_lines.append(f" - {metric.upper()}: {mean:.3f} ± {std:.3f}")
                cv_metrics_list.append({'Metric': metric.upper(), 'Mean': mean, 'Std': std})
        else:
            metrics = ['accuracy', 'precision', 'recall', 'f1']
            for metric in metrics:
                mean = cv_results[f'test_{metric}'].mean()
                std = cv_results[f'test_{metric}'].std()
                report_lines.append(f" - {metric.upper()}: {mean:.3f} ± {std:.3f}")
                cv_metrics_list.append({'Metric': metric.upper(), 'Mean': mean, 'Std': std})
            if 'test_roc_auc' in cv_results:
                mean = cv_results['test_roc_auc'].mean()
                std = cv_results['test_roc_auc'].std()
                report_lines.append(f" - ROC AUC: {mean:.3f} ± {std:.3f}")
                cv_metrics_list.append({'Metric': 'ROC AUC', 'Mean': mean, 'Std': std})
        cv_metrics_df = pd.DataFrame(cv_metrics_list)

        # Training metrics
        y_pred = self.model.predict(self.features)
        if hasattr(self, 'target_encoder'):
            y_pred = self.target_encoder.inverse_transform(y_pred)
            target = self.target_encoder.inverse_transform(self.target)
        else:
            target = self.target

        report_lines.append("\n📊 Training Set Metrics:")
        train_metrics_list = []
        if self.problem_type == "regression":
            train_metrics = {
                'R2': r2_score(target, y_pred),
                'MAE': mean_absolute_error(target, y_pred),
                'MSE': mean_squared_error(target, y_pred),
                'RMSE': np.sqrt(mean_squared_error(target, y_pred)),
            }
            for metric, value in train_metrics.items():
                report_lines.append(f" - {metric}: {value:.3f}")
                train_metrics_list.append({'Metric': metric, 'Value': value})
        else:
            train_metrics = {
                'Accuracy': accuracy_score(target, y_pred),
                'Precision': precision_score(target, y_pred, average='weighted', zero_division=0),
                'Recall': recall_score(target, y_pred, average='weighted', zero_division=0),
                'F1': f1_score(target, y_pred, average='weighted', zero_division=0)
            }
            if hasattr(self.model, 'predict_proba'):
                try:
                    train_metrics['ROC AUC'] = roc_auc_score(
                        target,
                        self.model.predict_proba(self.features),
                        multi_class='ovr'
                    )
                except:
                    pass
            for metric, value in train_metrics.items():
                report_lines.append(f" - {metric}: {value:.3f}")
                train_metrics_list.append({'Metric': metric, 'Value': value})
        train_metrics_df = pd.DataFrame(train_metrics_list)

        # Feature importance section
        report_lines, feature_df = _get_feature_importance_section(self)

        # Statistical details section
        if hasattr(self, 'stats_model') and self.p_values is not None:
            report_lines.extend(_get_statistical_details_section(self))

        # Export report
        if export_path:
            try:
                export_dir = os.path.dirname(os.path.abspath(export_path)) if os.path.dirname(export_path) else '.'
                os.makedirs(export_dir, exist_ok=True)
                
                if format.lower() == 'txt':
                    _export_txt_report(export_path, report_lines)
                    print(f"\n🦈 Text report exported to: {export_path}")
                elif format.lower() in ['pdf', 'docx']:
                    docx_path = export_path if format.lower() == 'docx' else export_path.replace('.pdf', '.docx')
                    _export_docx_report(docx_path, self, cv_metrics_df, train_metrics_df, self.problem_type)
                    if format.lower() == 'pdf':
                        print("⚠️ Note: PDF export requires Microsoft Word or LibreOffice installed.")
                        _convert_docx_to_pdf(docx_path, export_path)
                        os.remove(docx_path)
                        print(f"\n🦈 PDF report exported to: {export_path}")
                    else:
                        print(f"\n🦈 Word document exported to: {export_path}")
                else:
                    raise ValueError(f"🦈 Unsupported format: {format}")
                    
            except Exception as e:
                print(f"\n⚠️ Failed to export report: {str(e)}")
                raise

        # Print report to console
        for line in report_lines:
            print(line)

        return cv_results, train_metrics
    
    finally:
        # Restore original backend
        matplotlib.use(original_backend)