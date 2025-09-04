# sharkpy/explaining.py

import numpy as np
import pandas as pd
from textwrap import dedent
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert
import matplotlib
import os
import tempfile
from sklearn.metrics import (
    r2_score, mean_absolute_error, mean_squared_error,
    accuracy_score, precision_score, recall_score, f1_score,
    roc_auc_score, confusion_matrix
)
from sklearn.inspection import permutation_importance
from .plotting import plot_model
try:
    from .shapash_integration import explain_with_shapash
except ImportError:
    explain_with_shapash = None

def _create_temp_plot(model, features, target, kind: str, width: int = 8, height: int = 6) -> str:
    """Create a temporary plot and return its path"""
    try:
        fd, temp_path = tempfile.mkstemp(suffix='.png')
        os.close(fd)
        
        original_backend = matplotlib.get_backend()
        try:
            matplotlib.use('Agg')
            if isinstance(features, pd.DataFrame):
                feature_names = features.columns
            else:
                feature_names = [f"Feature_{i}" for i in range(features.shape[1])]
            plot_model(model, features, target, kind=kind, show=False, save_path=temp_path, feature_names=feature_names)
        finally:
            matplotlib.use(original_backend)
        
        return temp_path
    except Exception as e:
        print(f"⚠️ Could not create {kind} plot: {str(e)}")
        return None

def _add_table_to_doc(doc: Document, df: pd.DataFrame, title: str):
    """Helper function to add a pandas DataFrame as a table to DOCX"""
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Add header
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
    
    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = f"{value:.3f}" if isinstance(value, float) else str(value)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(12)

def _add_metrics_table_to_doc(doc: Document, cv_results, train_metrics):
    """Add performance metrics table to DOCX document"""
    doc.add_heading('Performance Metrics', level=1)
    
    is_regression = 'r2' in train_metrics
    metrics_data = [['Metric', 'Training', 'Cross-Validation']]
    if is_regression:
        metrics_data.extend([
            ['R² Score', f"{train_metrics['r2']:.3f}", f"{cv_results['test_r2'].mean():.3f} ± {cv_results['test_r2'].std():.3f}"],
            ['MAE', f"{train_metrics['mae']:.3f}", f"{cv_results['test_mae'].mean():.3f} ± {cv_results['test_mae'].std():.3f}"],
            ['MSE', f"{train_metrics['mse']:.3f}", f"{cv_results['test_mse'].mean():.3f} ± {cv_results['test_mse'].std():.3f}"],
            ['RMSE', f"{train_metrics['rmse']:.3f}", f"{cv_results['test_rmse'].mean():.3f} ± {cv_results['test_rmse'].std():.3f}"]
        ])
    else:
        metrics_data.extend([
            ['Accuracy', f"{train_metrics['accuracy']:.3f}", f"{cv_results['test_accuracy'].mean():.3f} ± {cv_results['test_accuracy'].std():.3f}"],
            ['Precision', f"{train_metrics['precision']:.3f}", f"{cv_results['test_precision'].mean():.3f} ± {cv_results['test_precision'].std():.3f}"],
            ['Recall', f"{train_metrics['recall']:.3f}", f"{cv_results['test_recall'].mean():.3f} ± {cv_results['test_recall'].std():.3f}"],
            ['F1-Score', f"{train_metrics['f1']:.3f}", f"{cv_results['test_f1'].mean():.3f} ± {cv_results['test_f1'].std():.3f}"]
        ])
        if 'test_roc_auc' in cv_results:
            metrics_data.append(['ROC AUC', f"{train_metrics.get('roc_auc', 'N/A'):.3f}", 
                                f"{cv_results['test_roc_auc'].mean():.3f} ± {cv_results['test_roc_auc'].std():.3f}"])
    
    table = doc.add_table(rows=len(metrics_data), cols=len(metrics_data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(metrics_data):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
            if i == 0:
                table.cell(i, j).paragraphs[0].runs[0].bold = True
            table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(11)

def _export_deep_explanation_docx(docx_path, deep_explanation, model, features, target, feature_df, depth='deep'):
    """Export explanation to DOCX with specified depth level"""
    doc = Document()
    doc.add_heading('🦈 SharkPy Model Explanation', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    depth_levels = {
        'simple': deep_explanation['simple'],
        'mechanics': deep_explanation['mechanics'],
        'interpretation': deep_explanation['interpretation'],
        'actionable': deep_explanation['actionable']
    }
    
    # Export only the specified depth or all for 'deep'
    if depth == 'deep':
        for level_name, level_content in deep_explanation.items():
            doc.add_heading(level_name.upper(), level=1)
            for line in level_content:
                doc.add_paragraph(line)
            doc.add_paragraph("")
    elif depth in depth_levels:
        doc.add_heading(depth.upper(), level=1)
        for line in depth_levels[depth]:
            doc.add_paragraph(line)
        doc.add_paragraph("")
    else:
        doc.add_paragraph(f"⚠️ Invalid depth level '{depth}'. Including all levels.")
        for level_name, level_content in deep_explanation.items():
            doc.add_heading(level_name.upper(), level=1)
            for line in level_content:
                doc.add_paragraph(line)
            doc.add_paragraph("")
    
    if feature_df is not None:
        _add_table_to_doc(doc, feature_df, 'Feature Importance')
        doc.add_paragraph("🧠 WHAT THIS MEANS:")
        if 'Coefficient' in feature_df.columns:
            doc.add_paragraph("• Positive coefficients increase predictions or make the class more likely")
            doc.add_paragraph("• Negative coefficients decrease predictions or make the class less likely")
            doc.add_paragraph("• Larger absolute values indicate stronger impact")
            doc.add_paragraph("• For classification, coefficients reflect log-odds; use exp(coefficient) for odds ratios")
        else:
            doc.add_paragraph("• Higher importance scores indicate features the model relies on most")
            doc.add_paragraph("• Percent shows relative contribution to the model’s decisions")
            doc.add_paragraph("• Focus on top-ranked features to understand what drives predictions")
        doc.add_paragraph(f"• Top feature: {feature_df.iloc[0]['Feature']} ({feature_df.iloc[0]['Absolute_Impact'] if 'Absolute_Impact' in feature_df.columns else feature_df.iloc[0]['Importance']:.3f})")
        doc.add_paragraph(f"• Bottom feature: {feature_df.iloc[-1]['Feature']} ({feature_df.iloc[-1]['Absolute_Impact'] if 'Absolute_Impact' in feature_df.columns else feature_df.iloc[-1]['Importance']:.3f})")
        doc.add_paragraph("• Consider focusing on high-impact features for better results")
    
    try:
        temp_path = _create_temp_plot(model, features, target, "feature_importance")
        if temp_path:
            doc.add_heading("Feature Importance Visualization", level=1)
            doc.add_picture(temp_path, width=Inches(5))
            os.remove(temp_path)
    except Exception as e:
        doc.add_paragraph(f"⚠️ Could not generate feature importance plot: {str(e)}")
    
    doc.save(docx_path)

def _convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using docx2pdf"""
    try:
        convert(docx_path, pdf_path)
    except Exception as e:
        print(f"⚠️ Failed to convert DOCX to PDF: {str(e)}")
        raise

def _get_data_insights(features, target, target_name, data, label_encoder, is_regression):
    """Extract dynamic insights from the data"""
    insights = {}
    
    # Feature-related insights
    if isinstance(features, pd.DataFrame):
        insights['num_features'] = features.shape[1]
        insights['feature_names'] = list(features.columns)
        # Get range or unique values for top feature
        top_feature = insights['feature_names'][0] if insights['feature_names'] else "unknown feature"
        if features[top_feature].dtype in ['float64', 'int64']:
            insights['top_feature_range'] = f"{features[top_feature].min():.2f} to {features[top_feature].max():.2f}"
        else:
            insights['top_feature_unique'] = len(features[top_feature].unique())
    else:
        insights['num_features'] = features.shape[1]
        insights['feature_names'] = [f"Feature_{i}" for i in range(features.shape[1])]
        insights['top_feature_range'] = "N/A"
    
    # Target-related insights
    if is_regression:
        insights['target_type'] = "numeric"
        insights['target_range'] = f"{target.min():.2f} to {target.max():.2f}"
        insights['target_mean'] = target.mean()
        insights['class_names'] = [target_name]  # Use target column name for regression
    else:
        insights['target_type'] = "categorical"
        insights['num_classes'] = len(np.unique(target))
        # Try to get original class names from label_encoder or data
        if label_encoder is not None and hasattr(label_encoder, 'classes_'):
            insights['class_names'] = list(label_encoder.classes_)
        elif data is not None and target_name in data.columns:
            insights['class_names'] = list(data[target_name].unique())
        else:
            insights['class_names'] = [str(x) for x in np.unique(target)]  # Fallback to encoded labels
        insights['class_distribution'] = pd.Series(target).value_counts(normalize=True).to_dict()
    
    return insights

def explain_model(model, features, target, target_name, data=None, label_encoder=None, cv_results=None, train_metrics=None, export_path=None, format='txt', depth='deep'):
    """Generate model explanation with customizable depth and export options"""
    print(f"🦈 Debug: Processing depth={depth}")
    if not hasattr(model, 'predict'):
        raise ValueError("🦈 Model must have a predict method!")
    
    report_lines = ["🦈 SHARKPY MODEL EXPLANATION", "=" * 60, ""]
    
    # Handle shapash case
    if depth == 'shapash':
        if explain_with_shapash is None:
            raise ImportError("🦈 Shapash is not installed. Install it with `pip install shapash`.")
        return explain_with_shapash(model, title_story=None, display=True)
    
    # Generate feature importance
    feature_data = []
    is_multi_class = False
    if isinstance(features, pd.DataFrame):
        feature_names = features.columns
    else:
        feature_names = [f"Feature_{i}" for i in range(features.shape[1])]
    
    if hasattr(model, 'feature_importances_'):
        importances = model.feature_importances_
        indices = np.argsort(importances)[::-1]
        for idx in indices:
            feature_data.append({
                'Feature': feature_names[idx],
                'Importance': importances[idx],
                'Absolute_Impact': importances[idx],
                'Rank': idx + 1
            })
    elif hasattr(model, 'coef_'):
        coefs = model.coef_
        if len(coefs.shape) > 1:  # Multi-class logistic regression
            is_multi_class = True
            for class_idx in range(coefs.shape[0]):
                indices = np.argsort(np.abs(coefs[class_idx]))[::-1]
                for idx in indices:
                    feature_data.append({
                        'Feature': feature_names[idx],
                        'Coefficient': coefs[class_idx][idx],
                        'Absolute_Impact': np.abs(coefs[class_idx][idx]),
                        'Rank': idx + 1,
                        'Class': class_idx
                    })
        else:
            indices = np.argsort(np.abs(coefs))[::-1]
            for idx in indices:
                feature_data.append({
                    'Feature': feature_names[idx],
                    'Coefficient': coefs[idx],
                    'Absolute_Impact': np.abs(coefs[idx]),
                    'Rank': idx + 1
                })
    else:
        try:
            perm_importance = permutation_importance(model, features, target, n_repeats=10, random_state=42)
            indices = np.argsort(perm_importance.importances_mean)[::-1]
            for idx in indices:
                feature_data.append({
                    'Feature': feature_names[idx],
                    'Importance': perm_importance.importances_mean[idx],
                    'Absolute_Impact': np.abs(perm_importance.importances_mean[idx]),
                    'Rank': idx + 1
                })
        except Exception:
            feature_data = None
    
    feature_df = pd.DataFrame(feature_data) if feature_data else None
    
    # Generate data insights
    model_type = type(model).__name__
    is_regression = model_type in ['LinearRegression', 'Ridge', 'Lasso', 'RandomForestRegressor', 
                                  'XGBRegressor', 'LGBMRegressor', 'CatBoostRegressor', 'SVR']
    data_insights = _get_data_insights(features, target, target_name, data, label_encoder, is_regression)
    
    # Generate dynamic explanations
    deep_explanation = {
        'simple': ["🎯 SIMPLE LEVEL:", "=" * 50],
        'mechanics': ["🔧 MECHANICS LEVEL:", "=" * 50],
        'interpretation': ["📊 INTERPRETATION LEVEL:", "=" * 50],
        'actionable': ["🚀 ACTIONABLE INSIGHTS:", "=" * 50]
    }
    
    # SIMPLE LEVEL
    if is_regression:
        deep_explanation['simple'].extend([
            f"This model predicts a number ranging from {data_insights['target_range']} (like {data_insights['target_mean']:.2f} on average).",
            f"It uses {data_insights['num_features']} features, with {feature_df.iloc[0]['Feature'] if feature_df is not None else 'a key feature'} leading the way!",
            "Think of it as a super-smart calculator crunching your data! 🧮"
        ])
    else:
        class_names = ', '.join(map(str, data_insights['class_names']))
        deep_explanation['simple'].extend([
            f"This model predicts one of {data_insights['num_classes']} categories ({class_names}).",
            f"It uses {data_insights['num_features']} features, with {feature_df.iloc[0]['Feature'] if feature_df is not None else 'a key feature'} being the star player!",
            "Imagine it drawing lines to sort your data into groups! 🎯"
        ])
    
    # MECHANICS LEVEL
    if model_type in ['LinearRegression', 'Ridge', 'Lasso']:
        top_coef = feature_df.iloc[0]['Coefficient'] if feature_df is not None and 'Coefficient' in feature_df.columns else 0
        deep_explanation['mechanics'].extend([
            f"This {model_type} model uses {data_insights['num_features']} features to make predictions.",
            f"It assigns weights to each feature, like {feature_df.iloc[0]['Feature'] if feature_df is not None else 'a key feature'} (weight: {top_coef:.3f}).",
            "It combines these weights in a straight-line formula to predict the outcome.",
            f"{'Ridge' if model_type == 'Ridge' else 'Lasso' if model_type == 'Lasso' else 'Linear'} uses {'L2' if model_type == 'Ridge' else 'L1' if model_type == 'Lasso' else 'no'} regularization to keep things stable."
        ])
    elif model_type in ['RandomForestRegressor', 'RandomForestClassifier']:
        deep_explanation['mechanics'].extend([
            f"This Random Forest combines many decision trees to predict {'numbers' if is_regression else 'categories'}.",
            f"It uses {data_insights['num_features']} features, with {feature_df.iloc[0]['Feature'] if feature_df is not None else 'a key feature'} being the most important.",
            "Each tree votes, and the model averages (for numbers) or picks the majority (for categories).",
            "It avoids overfitting by using random subsets of data and features."
        ])
    else:
        deep_explanation['mechanics'].extend([
            f"This {model_type} model uses {data_insights['num_features']} features to make predictions.",
            f"Top feature: {feature_df.iloc[0]['Feature'] if feature_df is not None else 'unknown'} (impact: {feature_df.iloc[0]['Absolute_Impact']:.3f} if feature_df is not None else 'N/A').",
            "It learns complex patterns in your data to make accurate predictions."
        ])
    
    # INTERPRETATION LEVEL
    if cv_results and train_metrics:
        if is_regression:
            baseline_r2 = 0  # Baseline: predicting the mean
            deep_explanation['interpretation'].extend([
                f"R² Score: {train_metrics['r2']:.3f} (training) vs. {cv_results['test_r2'].mean():.3f} ± {cv_results['test_r2'].std():.3f} (cross-validation).",
                f"This means the model explains {train_metrics['r2']*100:.1f}% of the variation in {data_insights['target_range']}, compared to 0% for a simple mean prediction.",
                f"MAE: {train_metrics['mae']:.3f} (training) vs. {cv_results['test_mae'].mean():.3f} ± {cv_results['test_mae'].std():.3f} (cross-validation).",
                "MAE shows the average prediction error in the same units as your target."
            ])
            if abs(train_metrics['r2'] - cv_results['test_r2'].mean()) > 0.1:
                deep_explanation['interpretation'].append("The gap between training and cross-validation suggests potential overfitting—watch out!")
        else:
            baseline_accuracy = 1 / data_insights['num_classes']  # Random guessing
            deep_explanation['interpretation'].extend([
                f"Accuracy: {train_metrics['accuracy']:.3f} (training) vs. {cv_results['test_accuracy'].mean():.3f} ± {cv_results['test_accuracy'].std():.3f} (cross-validation).",
                f"This is much better than random guessing ({baseline_accuracy*100:.1f}% for {data_insights['num_classes']} classes)!",
                f"F1-Score: {train_metrics['f1']:.3f} (training) vs. {cv_results['test_f1'].mean():.3f} ± {cv_results['test_f1'].std():.3f} (cross-validation).",
                "F1-Score balances precision and recall for reliable predictions."
            ])
            if abs(train_metrics['accuracy'] - cv_results['test_accuracy'].mean()) > 0.1:
                deep_explanation['interpretation'].append("The gap between training and cross-validation suggests potential overfitting—watch out!")
    
    # ACTIONABLE INSIGHTS
    if feature_df is not None:
        top_feature = feature_df.iloc[0]['Feature']
        low_feature = feature_df.iloc[-1]['Feature']
        deep_explanation['actionable'].extend([
            f"Focus on {top_feature}, as it has the biggest impact (score: {feature_df.iloc[0]['Absolute_Impact']:.3f}).",
            f"Consider dropping or re-evaluating {low_feature}, as it contributes the least (score: {feature_df.iloc[-1]['Absolute_Impact']:.3f}).",
            "Collect more data, especially for outliers or underrepresented cases.",
            f"Try other models using shark.battle() to see if you can beat the current performance!"
        ])
        if not is_regression and data_insights['class_distribution']:
            # Highlight imbalanced classes
            max_class = max(data_insights['class_distribution'], key=data_insights['class_distribution'].get)
            min_class = min(data_insights['class_distribution'], key=data_insights['class_distribution'].get)
            if data_insights['class_distribution'][max_class] / data_insights['class_distribution'][min_class] > 2:
                deep_explanation['actionable'].append(
                    f"Class '{data_insights['class_names'][int(max_class)] if max_class.isdigit() else max_class}' is much more common than '{data_insights['class_names'][int(min_class)] if min_class.isdigit() else min_class}'. Collect more data for '{data_insights['class_names'][int(min_class)] if min_class.isdigit() else min_class}' to balance the dataset."
                )
    else:
        deep_explanation['actionable'].extend([
            "Try feature engineering to create new features from your data.",
            "Collect more data to help the model learn better patterns.",
            f"Experiment with other models using shark.battle() to improve performance."
        ])
    
    depth_levels = {
        'simple': deep_explanation['simple'],
        'mechanics': deep_explanation['mechanics'],
        'interpretation': deep_explanation['interpretation'],
        'actionable': deep_explanation['actionable']
    }
    
    if depth == 'deep':
        for level_name, level_content in deep_explanation.items():
            report_lines.extend(level_content)
            report_lines.append("")
    elif depth in depth_levels:
        report_lines.extend(depth_levels[depth])
    else:
        report_lines.extend([f"⚠️ Invalid depth level '{depth}'. Using 'deep' instead.", ""])
        for level_content in deep_explanation.values():
            report_lines.extend(level_content)
            report_lines.append("")
    
    # Add feature importance
    if feature_df is not None:
        report_lines.extend(["", "📊 FEATURE IMPORTANCE ANALYSIS", "=" * 40])
        report_lines.append(feature_df.to_string(index=False))
        report_lines.append("")
        if 'Coefficient' in feature_df.columns:
            report_lines.extend([
                "🧠 WHAT THIS MEANS:",
                "• Positive coefficients increase predictions (regression) or make the class more likely (classification)",
                "• Negative coefficients decrease predictions or make the class less likely",
                "• Larger absolute values indicate stronger impact",
                "• For classification, coefficients reflect log-odds; use exp(coefficient) for odds ratios",
                ""
            ])
        else:
            report_lines.extend([
                "🧠 WHAT THIS MEANS:",
                "• Higher importance scores indicate features the model relies on most",
                "• Percent shows relative contribution to the model’s decisions",
                "• Focus on top-ranked features to understand what drives predictions",
                ""
            ])
        report_lines.extend([
            "💡 Feature Insights:",
            f"• Top feature: {feature_df.iloc[0]['Feature']} ({feature_df.iloc[0]['Absolute_Impact'] if 'Absolute_Impact' in feature_df.columns else feature_df.iloc[0]['Importance']:.3f})",
            f"• Bottom feature: {feature_df.iloc[-1]['Feature']} ({feature_df.iloc[-1]['Absolute_Impact'] if 'Absolute_Impact' in feature_df.columns else feature_df.iloc[-1]['Importance']:.3f})",
            "• Consider focusing on high-impact features for better results"
        ])
    
    # Add visualizations to console
    if feature_df is not None:
        try:
            report_lines.extend(["", "🦈 Generating feature importance visualization..."])
            print("\n🦈 Generating feature importance visualization...")
            plot_model(model, features, target, kind="feature_importance", show=True, feature_names=feature_names)
        except Exception as e:
            report_lines.append(f"🦈 Could not generate visualization: {str(e)}")
            print(f"🦈 Could not generate visualization: {str(e)}")
    
    report_lines.extend(["", "=" * 60, "🦈 Sharky says: Keep learning and experimenting! 🚀"])
    
    # Export if requested
    if export_path:
        try:
            export_dir = os.path.dirname(os.path.abspath(export_path)) if os.path.dirname(export_path) else '.'
            os.makedirs(export_dir, exist_ok=True)
            
            if format.lower() == 'txt':
                with open(export_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(report_lines))
                print(f"🦈 Explanation exported to: {export_path}")
            elif format.lower() in ['pdf', 'docx']:
                docx_path = export_path if format.lower() == 'docx' else export_path.replace('.pdf', '.docx')
                _export_deep_explanation_docx(docx_path, deep_explanation, model, features, target, feature_df, depth=depth)
                if cv_results and train_metrics:
                    doc = Document(docx_path)
                    _add_metrics_table_to_doc(doc, cv_results, train_metrics)
                    doc.save(docx_path)
                if format.lower() == 'pdf':
                    print("🦈 Converting to PDF... (this may take a moment)")
                    _convert_docx_to_pdf(docx_path, export_path)
                    os.remove(docx_path)
                    print(f"🦈 PDF explanation exported to: {export_path}")
                else:
                    print(f"🦈 Word explanation exported to: {export_path}")
            else:
                raise ValueError(f"🦈 Unsupported format: {format}")
        except Exception as e:
            print(f"⚠️ Failed to export explanation: {str(e)}")
            raise
    
    return feature_df

def demo_explanation():
    """Demo the explanation system"""
    print("🦈 SharkPy Deep Explanation System Demo")
    print("=" * 50)
    print("\nAvailable depth levels:")
    print("• 'simple' - Beginner-friendly overview")
    print("• 'mechanics' - Technical how-it-works")  
    print("• 'interpretation' - Performance analysis")
    print("• 'actionable' - Practical recommendations")
    print("• 'deep' - Complete analysis (default)")
    print("• 'shapash' - Interactive SHAP dashboard (requires shapash package)")
    print("\nUsage: shark.explain(depth='mechanics')")