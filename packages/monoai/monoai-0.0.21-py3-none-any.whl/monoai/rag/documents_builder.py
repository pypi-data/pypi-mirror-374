import os
import uuid
from typing import List, Dict, Tuple, Optional
#from monoai.tools.webscraping import WebScraping

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class DocumentsBuilder:
    """
    A utility class for building document collections from various sources.
    
    This class provides methods to extract text content from files and web pages,
    split the content into manageable chunks with configurable size and overlap,
    and prepare the data for storage in vector databases.
    
    The DocumentsBuilder is designed to work seamlessly with the RAG system,
    producing output that can be directly used with vector database operations.
    
    Features:
    - File-based document extraction with UTF-8 encoding support
    - Text string processing for in-memory content
    - Web scraping with multiple engine options (requests, tavily, selenium)
    - Word document extraction (.doc and .docx formats)
    - PDF document extraction with metadata
    - Multiple chunking strategies (word, sentence, paragraph, fixed, semantic)
    - Configurable chunk size and overlap parameters
    - Rich metadata generation for each document chunk
    - Unique ID generation for database storage
    
    Attributes:
        _chunk_strategy (str): The chunking strategy to use
        _chunk_size (int): Maximum size of each text chunk in characters
        _chunk_overlap (int): Number of characters to overlap between chunks
    
    Examples:
    --------
    Basic usage with file processing:
    
    ```python
    # Initialize with default chunk settings (word-based)
    builder = DocumentsBuilder()
    
    # Process a text file
    documents, metadatas, ids = builder.from_file("document.txt")
    
    # Add to vector database
    vector_db.add(documents, metadatas, ids)
    ```
    
    Text string processing:
    
    ```python
    # Process text strings directly
    text_content = "This is a long text that needs to be processed..."
    documents, metadatas, ids = builder.from_str(text_content)
    
    # Process with custom source name
    documents, metadatas, ids = builder.from_str(
        text_content,
        source_name="user_input"
    )
    ```
    
    Different chunking strategies:
    
    ```python
    # Default settings (word-based chunking)
    builder = DocumentsBuilder()
    
    # Sentence-based chunking (5 sentences per chunk)
    builder = DocumentsBuilder(chunk_strategy="sentence", chunk_size=5)
    
    # Paragraph-based chunking (3 paragraphs per chunk)
    builder = DocumentsBuilder(chunk_strategy="paragraph", chunk_size=3)
    
    # Fixed-size chunks (800 characters per chunk)
    builder = DocumentsBuilder(chunk_strategy="fixed", chunk_size=800)
    
    # Word-based chunks (50 words per chunk)
    builder = DocumentsBuilder(chunk_strategy="word", chunk_size=50)
    ```
    
    Web scraping with different engines:
    
    ```python
    # Basic web scraping
    documents, metadatas, ids = builder.from_url("https://example.com")
    
    # Advanced scraping with Tavily
    documents, metadatas, ids = builder.from_url(
        "https://example.com",
        engine="tavily",
        deep=True
    )
    
    # JavaScript-heavy sites with Selenium
    documents, metadatas, ids = builder.from_url(
        "https://spa-example.com",
        engine="selenium"
    )
    ```
    
    Word document processing:
    
    ```python
    # Process Word documents
    documents, metadatas, ids = builder.from_doc("document.docx")
    documents, metadatas, ids = builder.from_doc("document.doc")
    
    # Process with custom extraction method
    documents, metadatas, ids = builder.from_doc(
        "document.docx",
        extraction_method="docx2txt"
    )
    ```
    
    PDF document processing:
    
    ```python
    # Process PDF documents
    documents, metadatas, ids = builder.from_pdf("document.pdf")
    
    # Process with page range
    documents, metadatas, ids = builder.from_pdf(
        "document.pdf",
        page_range=(1, 10)  # Extract pages 1-10
    )
    ```
    
    Notes:
    ------
    - chunk_overlap should typically be 10-20% of chunk_size
    - chunk_overlap must be less than chunk_size to prevent infinite loops
    - Different strategies interpret chunk_size differently:
      * word: chunk_size = number of words per chunk
      * sentence: chunk_size = number of sentences per chunk
      * paragraph: chunk_size = number of paragraphs per chunk
      * fixed: chunk_size = number of characters per chunk
      * semantic: chunk_size = number of characters per chunk
    - Very small chunks may lose context
    - Very large chunks may be less focused for retrieval
    - Fixed and semantic strategies always produce chunks of exactly chunk_size (except the last one)
    - Word document processing requires python-docx and python-docx2txt packages
    - PDF processing requires PyPDF2 package
    """

    def __init__(self, chunk_strategy: str = "word", chunk_size: int = 1000, chunk_overlap: int = 0, custom_split_func: Optional[callable] = None):
        """
        Initialize the DocumentsBuilder with chunking parameters.
        
        Parameters:
        -----------
        chunk_strategy : str, default="word"
            The strategy to use for text chunking:
            - "word": Break at word boundaries (spaces and newlines) when possible
            - "sentence": Break at sentence boundaries (periods, exclamation marks, question marks)
            - "paragraph": Break at paragraph boundaries (double newlines)
            - "fixed": Break at exact character count without considering boundaries
            - "semantic": Break at semantic boundaries (headers, sections, etc.)
            - "custom": Use the provided custom_split_func for chunking
            
        chunk_size : int, default=1000
            The size limit for each chunk, interpreted differently based on strategy:
            - "word": Maximum number of words per chunk
            - "sentence": Maximum number of sentences per chunk  
            - "paragraph": Maximum number of paragraphs per chunk
            - "fixed": Maximum number of characters per chunk
            - "semantic": Maximum number of characters per chunk
            - "custom": Passed to custom_split_func as a parameter
            
        chunk_overlap : int, default=0
            The overlap between consecutive chunks, interpreted based on strategy:
            - "word": Number of words to overlap
            - "sentence": Number of sentences to overlap
            - "paragraph": Number of paragraphs to overlap
            - "fixed": Number of characters to overlap
            - "semantic": Number of characters to overlap
            - "custom": Passed to custom_split_func as a parameter
            
        custom_split_func : callable, optional
            Custom function to use for text splitting. If provided, automatically sets chunk_strategy to "custom"
            regardless of the chunk_strategy parameter value.
            The function should have the signature: func(text: str, chunk_size: int, chunk_overlap: int) -> List[str]
            and return a list of text chunks.
            
        Raises:
        -------
        ValueError
            If chunk_overlap >= chunk_size (would cause infinite loops)
            If chunk_size <= 0
            If chunk_overlap < 0
            If chunk_strategy="custom" but no custom_split_func is provided
            
        Examples:
        ---------
        ```python
        # Default settings (word-based chunking)
        builder = DocumentsBuilder()
        
        # Sentence-based chunking (5 sentences per chunk)
        builder = DocumentsBuilder(chunk_strategy="sentence", chunk_size=5)
        
        # Paragraph-based chunking (3 paragraphs per chunk)
        builder = DocumentsBuilder(chunk_strategy="paragraph", chunk_size=3)
        
        # Fixed-size chunks (800 characters per chunk)
        builder = DocumentsBuilder(chunk_strategy="fixed", chunk_size=800)
        
        # Word-based chunks (50 words per chunk)
        builder = DocumentsBuilder(chunk_strategy="word", chunk_size=50)
        
        # Custom chunking function
        def my_custom_split(text, chunk_size, chunk_overlap):
            # Split by lines and then by chunk_size
            lines = text.split('\n')
            chunks = []
            for i in range(0, len(lines), chunk_size - chunk_overlap):
                chunk_lines = lines[i:i + chunk_size]
                chunks.append('\n'.join(chunk_lines))
            return chunks
        
        # Strategy automatically set to "custom" when custom_split_func is provided
        builder = DocumentsBuilder(
            chunk_size=100,
            chunk_overlap=10,
            custom_split_func=my_custom_split
        )
        
        # Or explicitly set strategy (will be overridden to "custom")
        builder = DocumentsBuilder(
            chunk_strategy="word",  # This will be ignored
            chunk_size=100,
            chunk_overlap=10,
            custom_split_func=my_custom_split  # Strategy becomes "custom"
        )
        ```
        
        Notes:
        ------
        - chunk_overlap should typically be 10-20% of chunk_size
        - chunk_overlap must be less than chunk_size to prevent infinite loops
        - Different strategies interpret chunk_size differently:
          * word: chunk_size = number of words per chunk
          * sentence: chunk_size = number of sentences per chunk
          * paragraph: chunk_size = number of paragraphs per chunk
          * fixed: chunk_size = number of characters per chunk
          * semantic: chunk_size = number of characters per chunk
          * custom: chunk_size is passed to custom_split_func
        - Very small chunks may lose context
        - Very large chunks may be less focused for retrieval
        - Fixed and semantic strategies always produce chunks of exactly chunk_size (except the last one)
        - Custom functions should handle their own overlap logic
        - Custom functions can implement any splitting logic:
          * Split by specific delimiters (e.g., "---", "###")
          * Split by regex patterns
          * Split by semantic boundaries using NLP libraries
          * Split by document structure (headers, sections, etc.)
          * Combine multiple strategies
        - When custom_split_func is provided, chunk_strategy is automatically set to "custom"
          regardless of the chunk_strategy parameter value
        """
        # If custom_split_func is provided, automatically set strategy to "custom"
        if custom_split_func is not None:
            chunk_strategy = "custom"
        
        self._chunk_strategy = chunk_strategy
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._custom_split_func = custom_split_func
        
        # Validate parameters to prevent infinite loops
        if chunk_overlap >= chunk_size:
            raise ValueError(
                f"chunk_overlap ({chunk_overlap}) must be less than chunk_size ({chunk_size}) "
                "to prevent infinite loops. Recommended: chunk_overlap should be 10-20% of chunk_size."
            )
        
        if chunk_size <= 0:
            raise ValueError(f"chunk_size must be positive, got {chunk_size}")
        
        if chunk_overlap < 0:
            raise ValueError(f"chunk_overlap must be non-negative, got {chunk_overlap}")
        
        # Validate custom split function
        if chunk_strategy == "custom" and custom_split_func is None:
            raise ValueError("custom_split_func must be provided when chunk_strategy='custom'")
        
        if custom_split_func is not None and not callable(custom_split_func):
            raise ValueError("custom_split_func must be callable")

    def from_file(self, file_path: str) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Read a file and split it into chunks with specified size and overlap.
        
        This method reads a text file from the filesystem, splits its content
        into chunks according to the configured parameters, and generates
        metadata and unique IDs for each chunk.
        
        Parameters:
        -----------
        file_path : str
            Path to the text file to read. The file must exist and be
            readable. UTF-8 encoding is assumed.
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The text content split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              file information and chunk details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        FileNotFoundError
            If the specified file does not exist or is not accessible.
            
        UnicodeDecodeError
            If the file cannot be decoded as UTF-8.
            
        Examples:
        ---------
        ```python
        # Process a single file
        documents, metadatas, ids = builder.from_file("article.txt")
        
        # Access metadata information
        for i, metadata in enumerate(metadatas):
            print(f"Chunk {i+1}:")
            print(f"  File: {metadata['file_name']}")
            print(f"  Size: {metadata['chunk_size']} characters")
            print(f"  Position: {metadata['chunk_index'] + 1}/{metadata['total_chunks']}")
        ```
        
        Notes:
        ------
        - File is read entirely into memory before processing
        - Empty files will return empty lists
        - File path is stored in metadata for traceability
        - Chunk indexing starts at 0
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Read the file content
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_str(self, text: str, source_name: str = "text_string") -> Tuple[List[str], List[Dict], List[str]]:
        """
        Process a text string and split it into chunks with specified size and overlap.
        
        This method takes a text string directly and processes it using the same
        chunking logic as file processing. It's useful when you already have
        text content in memory and want to prepare it for vector database storage.
        
        Parameters:
        -----------
        text : str
            The text content to process and split into chunks.
            
        source_name : str, default="text_string"
            A descriptive name for the text source. This will be included
            in the metadata for traceability and identification.
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The text content split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              source information and chunk details
            - List of unique IDs: UUID strings for each chunk
            
        Examples:
        ---------
        ```python
        # Process a simple text string
        text_content = "This is a long text that needs to be processed..."
        documents, metadatas, ids = builder.from_str(text_content)
        
        # Process with custom source name
        documents, metadatas, ids = builder.from_str(
            text_content,
            source_name="user_input"
        )
        
        # Process multiple text strings
        text_parts = [
            "First part of the document...",
            "Second part of the document...",
            "Third part of the document..."
        ]
        
        all_documents = []
        all_metadatas = []
        all_ids = []
        
        for i, text_part in enumerate(text_parts):
            documents, metadatas, ids = builder.from_str(
                text_part,
                source_name=f"document_part_{i+1}"
            )
            all_documents.extend(documents)
            all_metadatas.extend(metadatas)
            all_ids.extend(ids)
        ```
        
        Notes:
        ------
        - Uses the same chunking strategy and parameters as other methods
        - Empty strings will return empty lists
        - Source name is stored in metadata for identification
        - Useful for processing text from APIs, user input, or generated content
        """
        if not text or not text.strip():
            return [], [], []
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'source_type': 'text_string',
                'source_name': source_name,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk),
                'chunk_strategy': self._chunk_strategy
            }
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_doc(self, file_path: str, extraction_method: str = "auto") -> Tuple[List[str], List[Dict], List[str]]:
        """
        Extract text from Word documents (.doc and .docx files) and split into chunks.
        
        This method supports both .doc and .docx formats using different extraction
        methods. For .docx files, it can use either python-docx or docx2txt libraries.
        For .doc files, it uses docx2txt which can handle the older format.
        
        Parameters:
        -----------
        file_path : str
            Path to the Word document (.doc or .docx file). The file must exist
            and be readable.
            
        extraction_method : str, default="auto"
            The method to use for text extraction:
            - "auto": Automatically choose the best method based on file extension
            - "docx": Use python-docx library (only for .docx files)
            - "docx2txt": Use docx2txt library (works for both .doc and .docx)
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The extracted text split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              file information, document properties, and chunk details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        FileNotFoundError
            If the specified file does not exist or is not accessible.
            
        ValueError
            If the file is not a supported Word document format or if the
            required extraction method is not available.
            
        ImportError
            If the required libraries for the chosen extraction method are not installed.
            
        Examples:
        ---------
        ```python
        # Process a .docx file with automatic method selection
        documents, metadatas, ids = builder.from_doc("document.docx")
        
        # Process a .doc file
        documents, metadatas, ids = builder.from_doc("document.doc")
        
        # Force specific extraction method
        documents, metadatas, ids = builder.from_doc(
            "document.docx",
            extraction_method="docx2txt"
        )
        
        # Access document metadata
        for metadata in metadatas:
            print(f"File: {metadata['file_name']}")
            print(f"Format: {metadata['document_format']}")
            print(f"Extraction method: {metadata['extraction_method']}")
        ```
        
        Notes:
        ------
        - .docx files are the modern Word format (Office 2007+)
        - .doc files are the legacy Word format (Office 97-2003)
        - python-docx provides better structure preservation for .docx files
        - docx2txt works with both formats but may lose some formatting
        - Document properties (title, author, etc.) are extracted when available
        - Images and complex formatting are not preserved in the extracted text
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file extension and validate
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension not in ['.doc', '.docx']:
            raise ValueError(f"Unsupported file format: {file_extension}. Only .doc and .docx files are supported.")
        
        # Determine extraction method
        if extraction_method == "auto":
            if file_extension == '.docx' and DOCX_AVAILABLE:
                extraction_method = "docx"
            else:
                extraction_method = "docx2txt"
        
        # Extract text based on method
        if extraction_method == "docx":
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx library is required for 'docx' extraction method. Install with: pip install python-docx")
            if file_extension != '.docx':
                raise ValueError("'docx' extraction method only supports .docx files")
            text, doc_properties = self._extract_with_docx(file_path)
        elif extraction_method == "docx2txt":
            if not DOCX2TXT_AVAILABLE:
                raise ImportError("docx2txt library is required for 'docx2txt' extraction method. Install with: pip install python-docx2txt")
            text, doc_properties = self._extract_with_docx2txt(file_path)
        else:
            raise ValueError(f"Unsupported extraction method: {extraction_method}")
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'document_format': file_extension[1:],  # Remove the dot
                'extraction_method': extraction_method,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            # Add document properties if available
            if doc_properties:
                metadata.update(doc_properties)
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_pdf(self, file_path: str, page_range: Optional[Tuple[int, int]] = None) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Extract text from PDF documents and split into chunks.
        
        This method extracts text content from PDF files using PyPDF2 library.
        It supports extracting all pages or a specific range of pages, and
        preserves page information in the metadata.
        
        Parameters:
        -----------
        file_path : str
            Path to the PDF file. The file must exist and be readable.
            
        page_range : Tuple[int, int], optional
            Range of pages to extract (start_page, end_page), where pages are
            1-indexed. If None, all pages are extracted.
            Example: (1, 5) extracts pages 1 through 5.
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The extracted text split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              file information, PDF properties, page information, and chunk details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        FileNotFoundError
            If the specified file does not exist or is not accessible.
            
        ValueError
            If the file is not a valid PDF or if the page range is invalid.
            
        ImportError
            If PyPDF2 library is not installed.
            
        Examples:
        ---------
        ```python
        # Process entire PDF
        documents, metadatas, ids = builder.from_pdf("document.pdf")
        
        # Process specific page range
        documents, metadatas, ids = builder.from_pdf(
            "document.pdf",
            page_range=(1, 10)  # Pages 1-10
        )
        
        # Process single page
        documents, metadatas, ids = builder.from_pdf(
            "document.pdf",
            page_range=(5, 5)  # Only page 5
        )
        
        # Access PDF metadata
        for metadata in metadatas:
            print(f"File: {metadata['file_name']}")
            print(f"Page: {metadata.get('page_number', 'N/A')}")
            print(f"Total pages: {metadata.get('total_pages', 'N/A')}")
            print(f"PDF title: {metadata.get('pdf_title', 'N/A')}")
        ```
        
        Notes:
        ------
        - Page numbers are 1-indexed (first page is page 1)
        - Text extraction quality depends on the PDF structure
        - Scanned PDFs may not extract text properly
        - PDF metadata (title, author, etc.) is extracted when available
        - Page information is preserved in chunk metadata
        - Images and complex formatting are not preserved
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 library is required for PDF processing. Install with: pip install PyPDF2")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Validate file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension != '.pdf':
            raise ValueError(f"Unsupported file format: {file_extension}. Only .pdf files are supported.")
        
        # Extract text and metadata from PDF
        text, pdf_properties, page_info = self._extract_from_pdf(file_path, page_range)
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'document_format': 'pdf',
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            # Add PDF properties if available
            if pdf_properties:
                metadata.update(pdf_properties)
            
            # Add page information if available
            if page_info:
                metadata.update(page_info)
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_url(self, url: str, engine: str = "requests", deep: bool = False) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Scrape content from a URL and split it into chunks with specified size and overlap.
        
        This method uses web scraping to extract text content from a webpage,
        then processes the content using the same chunking logic as file processing.
        Multiple scraping engines are supported for different types of websites.
        
        Parameters:
        -----------
        url : str
            The URL to scrape. Must be a valid HTTP/HTTPS URL.
            
        engine : str, default="requests"
            The web scraping engine to use:
            - "requests": Simple HTTP requests (fast, good for static content)
            - "tavily": Advanced web scraping with better content extraction
            - "selenium": Full browser automation (good for JavaScript-heavy sites)
            
        deep : bool, default=False
            If using the "tavily" engine, whether to use advanced extraction mode.
            Deep extraction provides better content quality but is slower.
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The scraped text split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              URL information and scraping details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        ValueError
            If the scraping fails or no text content is extracted.
            
        Examples:
        ---------
        ```python
        # Basic web scraping
        documents, metadatas, ids = builder.from_url("https://example.com")
        
        # Advanced scraping with Tavily
        documents, metadatas, ids = builder.from_url(
            "https://blog.example.com",
            engine="tavily",
            deep=True
        )
        
        # JavaScript-heavy site with Selenium
        documents, metadatas, ids = builder.from_url(
            "https://spa.example.com",
            engine="selenium"
        )
        
        # Access scraping metadata
        for metadata in metadatas:
            print(f"Source: {metadata['url']}")
            print(f"Engine: {metadata['scraping_engine']}")
            print(f"Deep extraction: {metadata['deep_extraction']}")
        ```
        
        Notes:
        ------
        - Scraping may take time depending on the engine and website complexity
        - Some websites may block automated scraping
        - Selenium requires Chrome/Chromium to be installed
        - Tavily requires an API key to be configured
        """

        result = scrape_web(url, engine=engine, deep=deep)
                
        if not result or not result.get("text"):
            raise ValueError(f"Failed to extract text content from URL: {url}")
        
        text = result["text"]
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'url': url,
                'source_type': 'web_page',
                'scraping_engine': engine,
                'deep_extraction': deep,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids
    
    def _extract_with_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from a .docx file using python-docx library.
        
        Parameters:
        -----------
        file_path : str
            Path to the .docx file
            
        Returns:
        --------
        Tuple[str, Dict]
            A tuple containing the extracted text and document properties
        """
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n\n".join(text_parts)
        
        # Extract document properties
        properties = {}
        core_props = doc.core_properties
        if core_props.title:
            properties['document_title'] = core_props.title
        if core_props.author:
            properties['document_author'] = core_props.author
        if core_props.subject:
            properties['document_subject'] = core_props.subject
        if core_props.created:
            properties['document_created'] = str(core_props.created)
        if core_props.modified:
            properties['document_modified'] = str(core_props.modified)
        
        return text, properties
    
    def _extract_with_docx2txt(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from a Word document (.doc or .docx) using docx2txt library.
        
        Parameters:
        -----------
        file_path : str
            Path to the Word document
            
        Returns:
        --------
        Tuple[str, Dict]
            A tuple containing the extracted text and document properties
        """
        text = docx2txt.process(file_path)
        
        # docx2txt doesn't provide document properties, so return empty dict
        properties = {}
        
        return text, properties
    
    def _extract_from_pdf(self, file_path: str, page_range: Optional[Tuple[int, int]] = None) -> Tuple[str, Dict, Dict]:
        """
        Extract text and metadata from a PDF file using PyPDF2.
        
        Parameters:
        -----------
        file_path : str
            Path to the PDF file
            
        page_range : Tuple[int, int], optional
            Range of pages to extract (start_page, end_page), 1-indexed
            
        Returns:
        --------
        Tuple[str, Dict, Dict]
            A tuple containing the extracted text, PDF properties, and page information
        """
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Get total number of pages
            total_pages = len(pdf_reader.pages)
            
            # Determine page range
            if page_range is None:
                start_page = 1
                end_page = total_pages
            else:
                start_page, end_page = page_range
                # Validate page range
                if start_page < 1 or end_page > total_pages or start_page > end_page:
                    raise ValueError(f"Invalid page range: {page_range}. Pages must be between 1 and {total_pages}")
            
            # Extract text from specified pages
            text_parts = []
            for page_num in range(start_page - 1, end_page):  # Convert to 0-indexed
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            
            # Extract PDF properties
            properties = {}
            if pdf_reader.metadata:
                metadata = pdf_reader.metadata
                if '/Title' in metadata:
                    properties['pdf_title'] = metadata['/Title']
                if '/Author' in metadata:
                    properties['pdf_author'] = metadata['/Author']
                if '/Subject' in metadata:
                    properties['pdf_subject'] = metadata['/Subject']
                if '/Creator' in metadata:
                    properties['pdf_creator'] = metadata['/Creator']
                if '/Producer' in metadata:
                    properties['pdf_producer'] = metadata['/Producer']
                if '/CreationDate' in metadata:
                    properties['pdf_creation_date'] = str(metadata['/CreationDate'])
                if '/ModDate' in metadata:
                    properties['pdf_modification_date'] = str(metadata['/ModDate'])
            
            # Add page information
            page_info = {
                'total_pages': total_pages,
                'extracted_pages_start': start_page,
                'extracted_pages_end': end_page,
                'extracted_pages_count': end_page - start_page + 1
            }
            
            return text, properties, page_info
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks using the specified chunking strategy.
        
        This private method implements different text chunking algorithms based
        on the configured chunk_strategy. It supports word, sentence, paragraph,
        fixed, and semantic chunking strategies.
        
        Parameters:
        -----------
        text : str
            The text content to split into chunks.
            
        Returns:
        --------
        List[str]
            List of text chunks based on the selected strategy.
            Empty chunks are automatically filtered out.
            
        Examples:
        ---------
        ```python
        # Internal usage (called by from_file, from_doc, from_pdf, and from_url)
        chunks = builder._split_text("This is a long text that needs to be split...")
        print(f"Created {len(chunks)} chunks using {builder._chunk_strategy} strategy")
        ```
        
        Notes:
        ------
        - Chunks are stripped of leading/trailing whitespace
        - Empty chunks are automatically filtered out
        - Different strategies have different characteristics:
          * word: Preserves word boundaries, good for general use
          * sentence: Preserves sentence context, good for Q&A
          * paragraph: Preserves paragraph context, good for document structure
          * fixed: Exact size control, may break words/sentences
          * semantic: Attempts to preserve semantic meaning
        """
        if len(text) <= self._chunk_size:
            return [text]
        
        if self._chunk_strategy == "word":
            return self._split_by_words(text)
        elif self._chunk_strategy == "sentence":
            return self._split_by_sentences(text)
        elif self._chunk_strategy == "paragraph":
            return self._split_by_paragraphs(text)
        elif self._chunk_strategy == "fixed":
            return self._split_fixed(text)
        elif self._chunk_strategy == "semantic":
            return self._split_semantic(text)
        elif self._chunk_strategy == "custom":
            return self._custom_split_func(text, self._chunk_size, self._chunk_overlap)
        else:
            raise ValueError(f"Unsupported chunk strategy: {self._chunk_strategy}")
    
    def _split_by_words(self, text: str) -> List[str]:
        """
        Split text by word boundaries while respecting word count.
        
        This strategy splits text into chunks based on the number of words,
        trying to break at word boundaries when possible.
        """
        # Split text into words
        words = text.split()
        
        if len(words) <= self._chunk_size:
            return [text]
        
        chunks = []
        start_word = 0
        
        while start_word < len(words):
            # Calculate end word position for current chunk
            end_word = start_word + self._chunk_size
            
            # Extract words for this chunk
            chunk_words = words[start_word:end_word]
            chunk = ' '.join(chunk_words)
            
            if chunk.strip():  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Calculate next start position with overlap
            new_start_word = end_word - self._chunk_overlap
            
            # Ensure we always advance to prevent infinite loops
            if new_start_word <= start_word:
                new_start_word = start_word + 1
            
            start_word = new_start_word
            
            # Safety check to prevent infinite loops
            if start_word >= len(words):
                break
        
        return chunks
    
    def _split_by_sentences(self, text: str) -> List[str]:
        """
        Split text by sentence boundaries while respecting sentence count.
        
        This strategy splits text into chunks based on the number of sentences,
        preserving sentence integrity.
        """
        # Define sentence endings
        sentence_endings = ['.', '!', '?', '\n\n']
        
        # Split text into sentences
        sentences = []
        last_pos = 0
        
        for i, char in enumerate(text):
            if char in sentence_endings:
                sentence = text[last_pos:i+1].strip()
                if sentence:
                    sentences.append(sentence)
                last_pos = i + 1
        
        # Add the last sentence if it doesn't end with punctuation
        if last_pos < len(text):
            last_sentence = text[last_pos:].strip()
            if last_sentence:
                sentences.append(last_sentence)
        
        if len(sentences) <= self._chunk_size:
            return [text]
        
        chunks = []
        start_sentence = 0
        
        while start_sentence < len(sentences):
            # Calculate end sentence position for current chunk
            end_sentence = start_sentence + self._chunk_size
            
            # Extract sentences for this chunk
            chunk_sentences = sentences[start_sentence:end_sentence]
            chunk = ' '.join(chunk_sentences)
            
            if chunk.strip():  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Calculate next start position with overlap
            new_start_sentence = end_sentence - self._chunk_overlap
            
            # Ensure we always advance to prevent infinite loops
            if new_start_sentence <= start_sentence:
                new_start_sentence = start_sentence + 1
            
            start_sentence = new_start_sentence
            
            # Safety check to prevent infinite loops
            if start_sentence >= len(sentences):
                break
        
        return chunks
    
    def _split_by_paragraphs(self, text: str) -> List[str]:
        """
        Split text by paragraph boundaries while respecting paragraph count.
        
        This strategy splits text into chunks based on the number of paragraphs,
        preserving paragraph integrity.
        """
        # Split by paragraph boundaries (double newlines)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if len(paragraphs) <= self._chunk_size:
            return [text]
        
        chunks = []
        start_paragraph = 0
        
        while start_paragraph < len(paragraphs):
            # Calculate end paragraph position for current chunk
            end_paragraph = start_paragraph + self._chunk_size
            
            # Extract paragraphs for this chunk
            chunk_paragraphs = paragraphs[start_paragraph:end_paragraph]
            chunk = '\n\n'.join(chunk_paragraphs)
            
            if chunk.strip():  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Calculate next start position with overlap
            new_start_paragraph = end_paragraph - self._chunk_overlap
            
            # Ensure we always advance to prevent infinite loops
            if new_start_paragraph <= start_paragraph:
                new_start_paragraph = start_paragraph + 1
            
            start_paragraph = new_start_paragraph
            
            # Safety check to prevent infinite loops
            if start_paragraph >= len(paragraphs):
                break
        
        return chunks
    
    def _split_fixed(self, text: str) -> List[str]:
        """
        Split text into fixed-size chunks without considering boundaries.
        
        This strategy creates chunks of exactly chunk_size characters
        (except possibly the last chunk) without trying to preserve
        word or sentence boundaries.
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self._chunk_size
            chunk = text[start:end].strip()
            
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Calculate next start position with overlap
            new_start = end - self._chunk_overlap
            
            # Ensure we always advance to prevent infinite loops
            if new_start <= start:
                new_start = start + 1
            
            start = new_start
            
            # Safety check to prevent infinite loops
            if start >= len(text):
                break
        
        return chunks
    
    def _split_semantic(self, text: str) -> List[str]:
        """
        Split text by semantic boundaries.
        
        This strategy attempts to break text at semantic boundaries like
        headers, section breaks, and other structural elements while
        respecting the chunk size.
        """
        # Define semantic break patterns
        semantic_patterns = [
            '\n# ', '\n## ', '\n### ', '\n#### ',  # Markdown headers
            '\n1. ', '\n2. ', '\n3. ', '\n4. ', '\n5. ',  # Numbered lists
            '\n• ', '\n- ', '\n* ',  # Bullet points
            '\n\n',  # Paragraph breaks
            '\n---\n', '\n___\n',  # Horizontal rules
            '\n\nChapter ', '\n\nSection ', '\n\nPart ',  # Document sections
        ]
        
        chunks = []
        current_chunk = ""
        
        # Split text by semantic patterns
        parts = [text]
        for pattern in semantic_patterns:
            new_parts = []
            for part in parts:
                if pattern in part:
                    split_parts = part.split(pattern)
                    for i, split_part in enumerate(split_parts):
                        if i > 0:  # Add the pattern back to all parts except the first
                            split_part = pattern + split_part
                        if split_part.strip():
                            new_parts.append(split_part)
                else:
                    new_parts.append(part)
            parts = new_parts
        
        # Group parts into chunks
        for part in parts:
            # If adding this part would exceed chunk size, start a new chunk
            if len(current_chunk) + len(part) > self._chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Start new chunk with overlap
                overlap_start = max(0, len(current_chunk) - self._chunk_overlap)
                current_chunk = current_chunk[overlap_start:] + part
            else:
                current_chunk += part
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
        